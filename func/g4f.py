from config import bot, g4f_client, get_client,Form, openai_clients
import config
from key import GEMINI_API_KEY
from aiogram.fsm.context import FSMContext
from database import load_context,save_context,def_enhance, def_gen_model, def_aspect
from aiogram import types
import asyncio
import logging
from io import BytesIO
import tempfile
import os
from func.openai_image import run_with_timeout
import aiohttp
from urllib.parse import quote
import random
from google import genai
from google.genai import types as genai_types
from datetime import timedelta
import time
from PIL import Image

new_api_models = ["flux", "turbo"]
fresed_models = ["stable-diffusion-3", "stable-diffusion-3-large", "stable-diffusion-3-large-turbo", "flux-pro-1.1", "flux-pro-1"]
google_ai_models = ["gemini-2.0-flash-exp"]

if GEMINI_API_KEY:
    genai_client = genai.Client(api_key=GEMINI_API_KEY)
else:
    genai_client = None
    logging.warning("GEMINI_API_KEY is not set. Google AI image generation will be unavailable.")

async def process_image_generation_prompt(message: types.Message, state: FSMContext):
    start_time = time.time()

    data = await state.get_data()
    prompt = data.get("image_generation_prompt")
    is_direct_image_gen = data.get("is_direct_image_gen", False)

    if not is_direct_image_gen:
        await message.reply(
            f"🔔Твой запрос '{message.text}' был переведен как '{prompt}'.\nНачинаю генерацию..."
        )

    await state.update_data(image_generation_prompt=prompt)

    user_id = message.from_user.id
    user_context = await load_context(user_id)
    DEFAULT_IMAGE_GEN_MODEL = await def_gen_model()
    DEFAULT_ASPECT_RATIO = await def_aspect()
    DEFAULT_ENHANCE = await def_enhance()
    
    model_info = user_context.get("image_generation_model", DEFAULT_IMAGE_GEN_MODEL)
    
    if isinstance(model_info, dict):
        model_id = model_info.get("model_id")
        api_type = model_info.get("api")
    else:
  
        last_underscore_pos = model_info.rfind('_')
        if last_underscore_pos != -1:
            model_id = model_info[:last_underscore_pos]
            api_type = model_info[last_underscore_pos+1:]
        else:
            model_id = model_info
            api_type = "poli"  # API по умолчанию
    
    aspect_ratio = user_context.get("aspect_ratio", DEFAULT_ASPECT_RATIO)
    enhance = user_context.get("enhance", DEFAULT_ENHANCE)
    
    await state.update_data(image_generation_model=model_id)
    await state.update_data(image_generation_model_api=api_type)
    await state.update_data(aspect_ratio=aspect_ratio)
    await state.update_data(enhance=enhance)

    aspect_ratio_options = {
        "1:1": (2048, 2048),
        "3:2": (1536, 1024),
        "2:3": (1024, 1536),
        "4:3": (1536, 1152),
        "3:4": (1152, 1536),
        "16:9": (1792, 1024),
        "9:16": (1024, 1792),
        "21:9": (2048, 896),
    }
    width, height = aspect_ratio_options.get(aspect_ratio, (1024, 1024)) 


    # Улучшение промпта для всех моделей, если enhance=True
    if enhance:
        try:
            system_prompt = "YOU ARE AN ELITE PROMPT ENGINEER SPECIALIZING IN ENHANCING PROMPTS FOR IMAGE DESCRIPTION GENERATION. YOUR TASK IS TO ACCEPT A TEXT INPUT IN ANY LANGUAGE AND PRODUCE AN OPTIMIZED PROMPT IN ENGLISH FOR A GENERATIVE MODEL. THE OPTIMIZED PROMPT MUST INCLUDE A DETAILED DESCRIPTION OF THE SCENE, SPECIFYING WHAT IS HAPPENING AND INCORPORATING SUBTLE DETAILS TO ENSURE BEAUTIFUL VISUALIZATION. YOUR OUTPUT SHOULD BE THE DESCRIPTION IN ENGLISH ONLY, WITHOUT ANY ADDITIONAL COMMENTS OR EXPLANATIONS. ###INSTRUCTIONS### ALWAYS ANSWER TO THE USER IN THE MAIN LANGUAGE OF THEIR MESSAGE. 1. **TRANSLATE** the provided input text to English if necessary. 2. **ANALYZE** the scene to identify all critical elements such as the setting, actions, characters, and objects. 3. **ENRICH** the description by adding subtle details that enhance the visual quality and realism of the scene. 4. **ENSURE** the prompt is clear. vivid. and evocative to aid the generative model in"
            improved_prompt = await asyncio.to_thread(
                lambda:
                    config.enhance_prompt_client.chat.completions.create(
                        model="llama-3.3-70b",
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": prompt}
                        ],
                    )
            )
            prompt = improved_prompt.choices[0].message.content
        except Exception as e:
            logging.error(f"Error during prompt improvement: {e}")
            await bot.send_message(user_id, f"🚨Ошибка при улучшении промпта: {e}")

    if api_type == "poli":
        async def fetch_image():
            encoded_prompt = quote(prompt)
            url = f"https://image.pollinations.ai/prompt/{encoded_prompt}"
            api_model_name = model_id 
            params = {
                "width": width,
                "height": height,
                "enhance": "false", 
                "model": api_model_name,
                "seed": random.randint(0, 1000000),
                "nologo": "true",
                "private": "true",
                "safe": "false"
            }
            try:
                async with aiohttp.ClientSession() as session:
                    async with session.get(url, params=params) as response:
                        if response.status == 200:
                            image_data = await response.read()
                            return image_data
                        else:
                            error_message = await response.text()
                            return response.status, error_message 
            except Exception as e:
                logging.error(f"Error during image generation: {e}")
                raise e
        retry_count = 0
        max_retries = 3
        while retry_count < max_retries:
            try:
                result = await asyncio.to_thread(lambda: asyncio.run(fetch_image()))
                if isinstance(result, tuple) and result[0] == 500: 
                    status_code, error_message = result
                    retry_count += 1
                    await bot.send_message(
                        user_id, f"⚠️ Ошибка 500 при генерации изображения. Попытка {retry_count}/{max_retries}..."
                    )
                    if retry_count == max_retries:
                        await bot.send_message(
                            user_id, f"🚨 Не удалось сгенерировать изображение после {max_retries} попыток. Ошибка: {error_message}"
                        )
                        break 
                    await asyncio.sleep(1) 
                else:
                    image_data = result
                    caption = f"Фото сгенерировано моделью {model_id}"
                    if aspect_ratio:
                        caption += f" с соотношением сторон {aspect_ratio}"
                    if enhance:
                        caption += f", enhance: {enhance}"
                    caption += ":"
                    await bot.send_photo(
                        user_id,
                        photo=types.BufferedInputFile(image_data, filename="image.jpg"),
                        caption=caption,
                    )
                    caption2 = "Фото без сжатия"
                    await bot.send_document(user_id, document=types.BufferedInputFile(image_data, filename="image.jpg"), caption=caption2)
                    break
            except Exception as e:
                await bot.send_message(
                    user_id, f"🚨Ошибка во время генерации изображения: {e}"
                )
                break 

    elif api_type in openai_clients and api_type != "poli":
        client = openai_clients.get(api_type)
        try:
            size_str = f"{width}x{height}"
            response = await run_with_timeout(
                client.images.generate(
                    model=model_id,
                    prompt=prompt,
                    size=size_str
                ),
                timeout=60,
                msg=message
            )
            if response is None:
                return
            if hasattr(response, 'generated_images'):
                image_data = response.generated_images[0].image.getvalue()
            elif hasattr(response, 'data') and response.data:
                image_url = response.data[0].url
                async with aiohttp.ClientSession() as session:
                    async with session.get(image_url) as resp:
                        if resp.status == 200:
                            image_data = await resp.read()
                        else:
                            raise Exception(f"Не удалось скачать изображение, статус: {resp.status}")
            else:
                raise Exception(f"Неподдерживаемый формат ответа от {api_type} client")
            caption = f"Фото сгенерировано {api_type} моделью {model_id}"
            if aspect_ratio:
                caption += f" с соотношением сторон {aspect_ratio}"
            if enhance:
                caption += f", enhance: {enhance}"
            caption += ":"
            await bot.send_photo(
                user_id,
                photo=types.BufferedInputFile(image_data, filename="image.jpg"),
                caption=caption,
            )
            caption2 = "Фото без сжатия"
            await bot.send_document(
                user_id,
                document=types.BufferedInputFile(image_data, filename="image.jpg"),
                caption=caption2,
            )
        except Exception as e:
            logging.error(f"Error during {api_type} image generation: {e}")
            await bot.send_message(user_id, f"🚨Ошибка во время генерации изображения с помощью {api_type} client: {e}")

    elif api_type == "gemini":
        if genai_client:
            try:
                response_coroutine = asyncio.to_thread(
                    lambda: genai_client.models.generate_content(
                        model=model_id,
                        contents=prompt,
                        config=genai_types.GenerateContentConfig(
                            response_modalities=['Text', 'Image']
                        )
                    )
                )
                
                response = await run_with_timeout(
                    response_coroutine,
                    timeout=60,
                    msg=message
                )
                
                if response is None:
                    return
                
                image_data = None
                text_response = ""
                
                if hasattr(response, 'candidates') and response.candidates and hasattr(response.candidates[0], 'content'):
                    for part in response.candidates[0].content.parts:
                        if part.text:
                            text_response += part.text
                        
                        if part.inline_data is not None and part.inline_data.mime_type.startswith('image/'):
                            image_data = part.inline_data.data
                
                if not image_data:
                    error_message = "🚨Модель не сгенерировала изображение в ответе."
                    if text_response:
                        error_message += f"\n\nОтвет модели: {text_response}"
                    await bot.send_message(user_id, error_message)
                    return

                caption = f"Фото сгенерировано моделью {model_id}"
                if enhance:
                    caption += f", enhance: {enhance}"
                caption += ":"
                
                await bot.send_photo(
                    user_id,
                    photo=types.BufferedInputFile(image_data, filename="image.jpg"),
                    caption=caption,
                )
                
                caption2 = "Фото без сжатия"
                await bot.send_document(
                    user_id, 
                    document=types.BufferedInputFile(image_data, filename="image.jpg"),
                    caption=caption2,
                )
            except Exception as e:
                logging.error(f"Error during Gemini image generation: {e}")
                await bot.send_message(user_id, f"🚨Ошибка во время генерации изображения с помощью Gemini API: {e}")
        else:
            await bot.send_message(user_id, "🚨Генерация изображений через Gemini недоступна. API-ключ не настроен.")

    elif api_type == "g4f":
        image_gen_client = get_client(user_id, "g4f_image_gen_client", model_name=model_id)
    
        try:
            response_coroutine = asyncio.to_thread(
                lambda: image_gen_client.images.generate(
                    prompt=prompt,
                    model=model_id,
                    response_format="url",
                    enhance=False,
                    private=True,
                    width=width,
                    height=height,
                )
            )
            response = await run_with_timeout(response_coroutine, timeout=60, msg=message)
    
            if response is None:
                return
            
            image_url = response.data[0].url
            caption = f"Фото сгенерировано моделью {model_id}"
            if aspect_ratio:
                caption += f" с соотношением сторон {aspect_ratio}"
            if enhance:
                caption += f", enhance: {enhance}"
            caption += ":"
            await bot.send_photo(
                user_id,
                photo=image_url,
                caption=caption,
            )
            caption2 = "Фото без сжатия"
            await bot.send_document(user_id, document=image_url, caption=caption2)
        except Exception as e:
            logging.error(f"Error during image generation: {e}")
            await bot.send_message(user_id, f"🚨Ошибка во время генерации изображения: {e}")

    end_time = time.time()
    processing_time = end_time - start_time
    formatted_processing_time = str(timedelta(seconds=int(processing_time)))
    service_info = f"⏳ Время обработки запроса: {formatted_processing_time}"

    if user_context.get("show_processing_time", True):
        await bot.send_message(user_id, service_info)

    await state.set_state(Form.waiting_for_message)
    await state.update_data(image_generation_prompt=None)
    await state.update_data(aspect_ratio=None)
    await state.update_data(enhance=None)


async def handle_image_recognition(message: types.Message, state: FSMContext):
    user_id = message.from_user.id

    photo = message.photo[-1]
    file_id = photo.file_id
    file = await bot.get_file(file_id)
    file_path = file.file_path
    image_data = await bot.download_file(file_path)
    image = BytesIO(image_data.read())
    image.seek(0)

    user_context = await load_context(user_id)
    user_context["g4f_image"] = image
    await save_context(user_id, user_context)

    await bot.send_message(user_id, "🔔Введите промпт для распознавания изображения:")


async def handle_files_or_urls(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    
    try:
        if message.document:
            processing_msg = await message.reply("🔔 Начинается обработка файла...")

            file_id = message.document.file_id
            file = await bot.get_file(file_id)
            file_path = file.file_path
            file_data = await bot.download_file(file_path)

            with tempfile.NamedTemporaryFile(
                delete=False, suffix=f"_{message.document.file_name}"
            ) as tmp_file:
                tmp_file.write(file_data.read())
                temp_file_path = tmp_file.name

            file_content = await asyncio.to_thread(process_local_file, temp_file_path)

            if file_content == "Unsupported file type":
                await processing_msg.edit_text(
                    "🚨 Неподдерживаемый тип файла. Поддерживаемые форматы:\n"
                    "- Документы: PDF, DOCX, DOC, XLSX, XLS\n"
                    "- Текстовые файлы: TXT, CSV, MD\n"
                    "- Код: PY, JS, PHP, HTML, XML, JSON, YAML, SQL и другие"
                )
                return
            elif file_content == "Error processing file":
                await processing_msg.edit_text(
                    "🚨 Произошла ошибка при обработке файла. Пожалуйста, попробуйте еще раз."
                )
                return

            user_context = await load_context(user_id)
            model_key = user_context["model"]  
            model_id, api_type = model_key.split('_')
            allowed_apis = list(openai_clients.keys()) + ["g4f"]

            if api_type == "gemini":
                last_message = user_context["messages"][-1] if user_context["messages"] else None
                if last_message and last_message["role"] == "user" and any("data" in part for part in last_message["parts"]):
                    user_context["messages"][-1]["parts"].append({"text": file_content})
                else:
                    user_context["messages"].append({"role": "user", "parts": [{"text": file_content}]})
            elif api_type in  allowed_apis:
                user_context["messages"].append({"role": "user", "content": file_content})

            await save_context(user_id, user_context)

            await processing_msg.edit_text("🔔 Файл успешно обработан и добавлен в контекст.")
            await state.set_state(Form.waiting_for_message)

    except Exception as e:
        logging.error(f"Ошибка при обработке файла: {e}")
        if 'processing_msg' in locals():
            await processing_msg.edit_text(f"🚨 Произошла ошибка при обработке файла: {e}")
        else:
            await message.reply(f"🚨 Произошла ошибка при обработке файла: {e}")
        await state.set_state(Form.waiting_for_message)
    finally:
        if "temp_file_path" in locals() and os.path.exists(temp_file_path):
            os.remove(temp_file_path)

def process_local_file(file_path):
    import fitz 
    import os
    import logging
    
    file_content = ""
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # PDF файлы
        if file_ext == ".pdf":
            with fitz.open(file_path) as doc:
                for page in doc:
                    file_content += page.get_text()
        
        # Microsoft Word (.docx) документы
        elif file_ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            for para in doc.paragraphs:
                file_content += para.text + "\n"
            # Также получаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    file_content += " | ".join(row_text) + "\n"
        
        # Старые Microsoft Word (.doc) документы
        elif file_ext == ".doc":
            try:
                # Используем antiword как основной способ
                import subprocess
                result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
                if result.returncode == 0:
                    file_content = result.stdout
                else:
                    raise Exception(f"antiword завершился с ошибкой: {result.stderr}")
            except Exception as e2:
                logging.warning(f"Не удалось обработать .doc с помощью antiword: {e2}")
                try:
                    # Пробуем через libreoffice как запасной вариант
                    import os
                    tmp_txt = f"{file_path}.txt"
                    result = subprocess.run(['libreoffice', '--headless', '--convert-to', 'txt', file_path, 
                                             '--outdir', os.path.dirname(file_path)], 
                                            capture_output=True, text=True)
                    
                    # Определяем имя выходного файла
                    base_name = os.path.basename(file_path)
                    file_name_without_ext = os.path.splitext(base_name)[0]
                    converted_txt = os.path.join(os.path.dirname(file_path), f"{file_name_without_ext}.txt")
                    
                    if os.path.exists(converted_txt):
                        with open(converted_txt, 'r', encoding='utf-8', errors='ignore') as f:
                            file_content = f.read()
                        os.remove(converted_txt)
                    else:
                        raise Exception("Конвертация не удалась")
                except Exception as e3:
                    logging.error(f"Все методы обработки .doc не удались: {e3}")
                    return "Error processing .doc file: All methods failed"
        
        # Microsoft Excel (.xlsx) таблицы
        elif file_ext == ".xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet in wb.worksheets:
                file_content += f"Лист: {sheet.title}\n"
                for row in sheet.iter_rows(values_only=True):
                    file_content += " | ".join([str(cell) if cell is not None else "" for cell in row]) + "\n"
                file_content += "\n"
        
        # Старые Microsoft Excel (.xls) таблицы
        elif file_ext == ".xls":
            try:
                import xlrd
                wb = xlrd.open_workbook(file_path)
                for sheet_index in range(wb.nsheets):
                    sheet = wb.sheet_by_index(sheet_index)
                    file_content += f"Лист: {sheet.name}\n"
                    for row_index in range(sheet.nrows):
                        row_values = sheet.row_values(row_index)
                        file_content += " | ".join([str(cell) if cell else "" for cell in row_values]) + "\n"
                    file_content += "\n"
            except Exception as e:
                logging.error(f"Ошибка при обработке .xls файла: {e}")
                try:
                    # Резервный метод через libreoffice
                    tmp_csv = f"{file_path}.csv"
                    result = subprocess.run(['libreoffice', '--headless', '--convert-to', 'csv', file_path, 
                                             '--outdir', os.path.dirname(file_path)], 
                                            capture_output=True, text=True)
                    if os.path.exists(tmp_csv):
                        with open(tmp_csv, 'r', encoding='utf-8', errors='ignore') as f:
                            file_content = f.read()
                        os.remove(tmp_csv)
                    else:
                        raise Exception("Конвертация не удалась")
                except Exception as e2:
                    logging.error(f"Все методы обработки .xls не удались: {e2}")
                    return "Error processing .xls file: All methods failed"
        
        # Стандартные текстовые файлы
        elif file_ext in (
            ".txt", ".xml", ".json", ".js", ".har", ".sh", ".py",
            ".php", ".css", ".yaml", ".sql", ".log", ".csv", ".twig", ".md",
            ".c", ".cpp", ".h", ".java", ".rb", ".pl", ".rs", ".go", ".ts", ".jsx", ".tsx",
            ".conf", ".ini", ".toml", ".lua", ".bat", ".ps1", ".yml"
        ):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                file_content += f.read()
        else:
            return "Unsupported file type"
            
        return file_content
    except Exception as e:
        logging.error(f"Ошибка при обработке файла {file_path}: {e}")
        return f"Error processing file: {str(e)}"

async def process_image_editing(message: types.Message, state: FSMContext):
    start_time = time.time()
    user_id = message.from_user.id
    
    data = await state.get_data()
    image_data = data.get("image_edit_data")
    instructions = data.get("image_edit_instructions")
    
    if not image_data or not instructions:
        await message.reply("🚨 Не удалось получить данные изображения или инструкции по редактированию")
        await state.set_state(Form.waiting_for_message)
        return
    
    user_context = await load_context(user_id)
    model_info = user_context.get("image_generation_model")
    
    if isinstance(model_info, dict):
        model_id = model_info.get("model_id")
        api_type = model_info.get("api")
    else:
        last_underscore_pos = model_info.rfind('_')
        if last_underscore_pos != -1:
            model_id = model_info[:last_underscore_pos]
            api_type = model_info[last_underscore_pos+1:]
        else:
            model_id = model_info
            api_type = "poli" 
    
    if api_type == "gemini" and model_id in google_ai_models:
        try:
            if isinstance(image_data, list):
                pil_images = []
                for img in image_data:
                    if hasattr(img, 'read'):
                        image_bytes = BytesIO(img.read())
                    else:
                        image_bytes = BytesIO(img)
                    image_bytes.seek(0)
                    pil_img = await asyncio.to_thread(lambda: Image.open(image_bytes))
                    pil_images.append(pil_img)
            else:
                image_bytes = BytesIO(image_data.read())
                image_bytes.seek(0)
                pil_img = await asyncio.to_thread(lambda: Image.open(image_bytes))
                pil_images = [pil_img]
    
            contents = [instructions, *pil_images]
    
            response_coroutine = asyncio.to_thread(
                lambda: genai_client.models.generate_content(
                    model=model_id,
                    contents=contents,
                    config=genai_types.GenerateContentConfig(
                        response_modalities=['Text', 'Image']
                    )
                )
            )
            
            response = await run_with_timeout(
                response_coroutine,
                timeout=60, 
                msg=message
            )
            
            if response is None:
                await bot.send_message(
                    user_id, 
                    "🚨 Превышено время ожидания при редактировании изображения. Пожалуйста, попробуйте еще раз."
                )
                await state.set_state(Form.waiting_for_message)
                return
            
            edited_image_data = None
            text_response = ""
            
            if hasattr(response, 'candidates') and response.candidates and hasattr(response.candidates[0], 'content'):
                for part in response.candidates[0].content.parts:
                    if hasattr(part, 'text') and part.text:
                        text_response += part.text
                    
                    if hasattr(part, 'inline_data') and part.inline_data and part.inline_data.mime_type.startswith('image/'):
                        edited_image_data = part.inline_data.data
            
            if not edited_image_data:
                error_message = "🚨 Модель не сгенерировала отредактированное изображение в ответе."
                if text_response:
                    error_message += "\n\nОтвет модели: " + text_response[:3000]
                await bot.send_message(user_id, error_message)
                await state.set_state(Form.waiting_for_message)
                return
            
            caption = f"✏️ Изображение отредактировано с помощью {model_id}"
            await bot.send_photo(
                user_id,
                photo=types.BufferedInputFile(edited_image_data, filename="edited_image.jpg"),
                caption=caption,
            )
            
            if text_response:
                await bot.send_message(user_id, text_response)
            
            await bot.send_document(
                user_id, 
                document=types.BufferedInputFile(edited_image_data, filename="edited_image.jpg"), 
                caption="Отредактированное изображение без сжатия"
            )
            
        except Exception as e:
            logging.error(f"Error during image editing: {e}")
            error_message = f"🚨 Ошибка при редактировании изображения: {str(e)[:200]}"
            await bot.send_message(user_id, error_message)
    else:
        await bot.send_message(
            user_id, 
            "🚨 Редактирование изображений поддерживается только моделями Google AI. Пожалуйста, измените модель в настройках."
        )
    
    end_time = time.time()
    processing_time = end_time - start_time
    formatted_processing_time = str(timedelta(seconds=int(processing_time)))
    
    if user_context.get("show_processing_time", True):
        await bot.send_message(user_id, f"⏳ Время обработки запроса: {formatted_processing_time}")
    
    await state.set_state(Form.waiting_for_message)
    await state.update_data(image_edit_data=None)
    await state.update_data(image_edit_instructions=None)
